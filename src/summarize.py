"""Reader summaries — two LLM stages plus a deterministic binder.

The client compiles a reader-facing publication of one-page-per-FIRM outlook
summaries. This module produces the raw Word document that feeds it, in three
steps, because a firm's sources can straddle runs and batches:

1. `digest`    — Stage 1, one LLM call per source (run-time). The ONLY stage
                 that reads documents: it opens the native document, is grounded
                 by that source's kept calls + rolling memory, and emits a
                 structured per-source digest JSON.
2. `firmpages` — Stage 2, one LLM call per firm (batch-combine time). Reads the
                 firm's digests + its RECONCILED final calls (never the original
                 documents), writes a one-page markdown summary.
3. `bind`      — deterministic `python-docx` merge of the firm pages into one
                 Word file (no LLM).

Between stages 1 and 2 sits Task 2's deterministic reconcile: N run outputs +
the cross-check verdicts over them -> one final call set per firm. It is the
v1 stopgap for the v1.2 dual-confidence firm-reconcile stage; `needs_human`
conflicts never block — they flow through as an in-page divergence note.

    .venv/bin/python -m src.summarize digest --run runs/<name> --out-dir <dir>
    .venv/bin/python -m src.summarize firmpages --digests <dir> \\
        --outputs <output.csv> [...] --crosscheck <crosscheck.csv> --out-dir <dir>
    .venv/bin/python -m src.summarize bind --pages <dir> --out <file.docx>

Nothing here ever writes under `runs/` — every command takes an explicit output
directory or file.
"""

from __future__ import annotations

import argparse
import csv
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

from src import llm
from src.crosscheck import Row, load_rows
from src.eval import _leaf_key, normalize_firm
from src.ingest import slugify

PROJECT_ROOT = Path(__file__).resolve().parents[1]
DIGEST_PROMPT = PROJECT_ROOT / "prompts" / "summarize_digest.md"
FIRM_PAGE_PROMPT = PROJECT_ROOT / "prompts" / "summarize_firm_page.md"

# Stage defaults (overridable by flags on every command, like every pipeline
# step). Stage 1 is a high-volume reading/extraction job -> codex; stage 2 is
# digest synthesis plus editorial prose -> claude/sonnet at high effort.
DIGEST_ENGINE, DIGEST_MODEL, DIGEST_EFFORT = "codex", None, "medium"
FIRMPAGE_ENGINE, FIRMPAGE_MODEL, FIRMPAGE_EFFORT = "claude", "sonnet", "high"

Runner = Callable[[list[str], str], object]

# The em-dash-delimited memory header written by run.py `_memory_header`:
# ``# {firm} — {title}  ({source_id})``.
_MEMORY_HEADER = re.compile(r"^#\s+(?P<firm>.+?)\s+—\s+(?P<title>.+?)\s+\((?P<sid>[^()]+)\)\s*$")

_PIPE_SPLIT = re.compile(r"\s*\|\s*")


class SummarizeError(RuntimeError):
    """A fatal problem locating or loading run artifacts."""


# --------------------------------------------------------------------------- #
# Task 1 plumbing — map a frozen run's artifacts back to per-source inputs
# --------------------------------------------------------------------------- #


@dataclass(frozen=True, slots=True)
class RunSource:
    """One source of a frozen run, with the material a digest call needs."""

    source_id: str
    firm: str
    title: str
    url: str
    date: str
    native_doc: Path
    memory_text: str
    kept_rows: tuple[dict[str, str], ...]


def _work_dir_for(run_dir: Path) -> Path:
    """The run's work directory holds the per-source ingest artifacts + memory.

    A default run writes outputs to ``runs/<id>`` and work to ``work/<id>``;
    an ``--out-root`` run writes ``<out-root>/<id>`` and ``<out-root>/work/<id>``
    (work is a sibling of the run dir). Prefer the sibling when it exists so
    out-root runs resolve without symlink workarounds; a default-layout run has
    no ``runs/work/<id>`` sibling, so it falls back to ``PROJECT_ROOT/work/<id>``."""
    sibling = run_dir.resolve().parent / "work" / run_dir.name
    if sibling.is_dir():
        return sibling
    return PROJECT_ROOT / "work" / run_dir.name


def _read_output_rows(run_dir: Path) -> list[dict[str, str]]:
    path = run_dir / "output.csv"
    if not path.is_file():
        raise SummarizeError(f"run output not found: {path}")
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def _aligned_segment(row_value: str, titles: list[str], title: str) -> str:
    """Pick this source's segment out of a pipe-joined combined-row field.

    Grouped sources are written as one combined row with `Source`/`Date`/`URL`
    pipe-joined in the same order. Given the row's `Source` segments and this
    source's title, return the matching segment of `row_value`; fall back to the
    whole value when the field is not pipe-joined the same way."""
    segments = _PIPE_SPLIT.split(row_value.strip())
    if len(segments) == len(titles) and title in titles:
        return segments[titles.index(title)].strip()
    return row_value.strip()


def _rows_for_source(rows: list[dict[str, str]], title: str) -> tuple[tuple[dict[str, str], ...], str, str]:
    """Kept rows for one source, plus its url/date drawn from those rows.

    A row belongs to the source when the source's title is one of the row's
    pipe-split `Source` segments (so a grouped combined row is attributed to
    both of its documents). url/date are taken from the first matching row,
    de-pipe-joined to this source's own segment."""
    matched: list[dict[str, str]] = []
    url = date = ""
    for row in rows:
        titles = [seg.strip() for seg in _PIPE_SPLIT.split((row.get("Source") or "").strip())]
        if title in titles:
            matched.append(row)
            if not url and not date:
                url = _aligned_segment(row.get("URL") or "", titles, title)
                date = _aligned_segment(row.get("Date") or "", titles, title)
    return tuple(matched), url, date


def load_run_sources(run_dir: Path, *, only: set[str] | None = None) -> list[RunSource]:
    """Enumerate the run's sources from its own artifacts (work dir + output).

    For each source directory under ``work/<id>`` that carries a `memory.md`:
    parse firm/title/source_id from the memory header, locate the native
    document via `chunks.json`, and attach that source's kept output rows (and
    its url/date). `only` filters to a subset of source ids. Sorted by
    source_id for deterministic ordering."""
    work_dir = _work_dir_for(run_dir)
    if not work_dir.is_dir():
        raise SummarizeError(f"run work directory not found: {work_dir}")
    output_rows = _read_output_rows(run_dir)

    sources: list[RunSource] = []
    for source_dir in sorted(p for p in work_dir.iterdir() if p.is_dir()):
        memory_path = source_dir / "memory.md"
        chunks_path = source_dir / "chunks.json"
        if not memory_path.is_file() or not chunks_path.is_file():
            continue
        if only is not None and source_dir.name not in only:
            continue
        memory_text = memory_path.read_text(encoding="utf-8")
        header = _MEMORY_HEADER.match(memory_text.splitlines()[0] if memory_text else "")
        if header is None:
            raise SummarizeError(f"unparseable memory header in {memory_path}")
        firm, title = header.group("firm").strip(), header.group("title").strip()
        native_doc = _native_doc_path(source_dir, chunks_path)
        kept_rows, url, date = _rows_for_source(output_rows, title)
        sources.append(
            RunSource(
                source_id=source_dir.name,
                firm=firm,
                title=title,
                url=url,
                date=date,
                native_doc=native_doc,
                memory_text=memory_text,
                kept_rows=kept_rows,
            )
        )
    if only is not None:
        missing = only - {s.source_id for s in sources}
        if missing:
            raise SummarizeError(f"requested source ids not found in run: {sorted(missing)}")
    return sources


def _native_doc_path(source_dir: Path, chunks_path: Path) -> Path:
    """The native document the run read, resolved next to the source dir.

    `chunks.json` records the source_path the run analyzed (a PDF, a printed
    capture, or the extracted-text snapshot). We resolve by basename against the
    source dir so it survives the repo moving on disk."""
    chunks = json.loads(chunks_path.read_text(encoding="utf-8"))
    if not chunks:
        raise SummarizeError(f"no chunks recorded in {chunks_path}")
    native = source_dir / Path(chunks[0]["source_path"]).name
    if not native.is_file():
        raise SummarizeError(f"native document missing for source: {native}")
    return native


# --------------------------------------------------------------------------- #
# Task 1 — Stage 1 per-source digest
# --------------------------------------------------------------------------- #

_STANCE_WORDS = frozenset({"overweight", "neutral", "underweight", "uncertain", "mixed"})


def parse_digest(raw_response: str) -> dict:
    """Validate the stage-1 digest contract; raise so the repair loop re-prompts."""
    payload = json.loads(llm._extract_json(raw_response))
    if not isinstance(payload, dict):
        raise ValueError("digest response must be a JSON object")
    for field in ("firm", "document_title"):
        value = payload.get(field)
        if not isinstance(value, str) or not value.strip():
            raise ValueError(f"digest needs a non-empty {field}")
    for field in ("url", "date"):
        if not isinstance(payload.get(field, ""), str):
            raise ValueError(f"digest {field} must be a string")
    themes = payload.get("themes")
    if not isinstance(themes, list):
        raise ValueError("digest themes must be a list")
    for theme in themes:
        if not isinstance(theme, dict):
            raise ValueError("each theme must be an object")
        if not isinstance(theme.get("label"), str) or not theme["label"].strip():
            raise ValueError("each theme needs a non-empty label")
        if not isinstance(theme.get("summary"), str):
            raise ValueError("each theme needs a summary string")
        points = theme.get("points", [])
        if not isinstance(points, list) or not all(isinstance(p, str) for p in points):
            raise ValueError("theme points must be a list of strings")
    stances = payload.get("stances")
    if not isinstance(stances, list):
        raise ValueError("digest stances must be a list")
    for stance in stances:
        if not isinstance(stance, dict):
            raise ValueError("each stance must be an object")
        if not isinstance(stance.get("asset_class"), str) or not stance["asset_class"].strip():
            raise ValueError("each stance needs a non-empty asset_class")
        if stance.get("stance") not in _STANCE_WORDS:
            raise ValueError(f"stance must be one of {sorted(_STANCE_WORDS)}; got {stance.get('stance')!r}")
        if not isinstance(stance.get("detail", ""), str):
            raise ValueError("stance detail must be a string")
    return payload


def digest_source(
    source: RunSource,
    *,
    engine: str,
    model: str | None,
    effort: str | None,
    runner: Runner | None = None,
) -> dict:
    """One LLM call: read the native document (grounded by kept calls + memory),
    emit a structured digest. The native document reaches the model the same way
    the checker's visual route attaches it — an absolute path in the inputs the
    prompt instructs the model to open."""
    inputs = {
        "source_id": source.source_id,
        "firm": source.firm,
        "document_title": source.title,
        "url": source.url,
        "date": source.date,
        "native_source_path": str(source.native_doc.resolve()),
        "kept_calls": [
            {
                "sub_asset_class": row.get("Sub-Asset Class", ""),
                "view": row.get("View", ""),
                "basis": row.get("basis", ""),
                "full_commentary": row.get("Full Commentary", ""),
            }
            for row in source.kept_rows
        ],
    }
    result = llm.call_parsed(
        DIGEST_PROMPT,
        inputs,
        engine=engine,
        model=model,
        effort=effort,
        runner=runner,
        template_vars={"memory": source.memory_text},
        parser=parse_digest,
    )
    return result.payload


def run_digests(
    run_dir: Path,
    out_dir: Path,
    *,
    only: set[str] | None = None,
    engine: str = DIGEST_ENGINE,
    model: str | None = DIGEST_MODEL,
    effort: str | None = DIGEST_EFFORT,
    runner: Runner | None = None,
) -> list[Path]:
    """Digest every selected source of the run into one JSON file per source."""
    sources = load_run_sources(run_dir, only=only)
    out_dir.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    for source in sources:
        payload = digest_source(
            source, engine=engine, model=model, effort=effort, runner=runner
        )
        path = out_dir / f"{source.source_id}.json"
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True), encoding="utf-8")
        written.append(path)
    return written


# --------------------------------------------------------------------------- #
# Task 2 — deterministic reconcile (pure code, no LLM)
# --------------------------------------------------------------------------- #

RES_SINGLE = "single"
RES_SAME_VIEW = "duplicate_same_view"
RES_SUPERSEDED = "superseded"
RES_SAME_CALL = "same_call"
RES_UNRESOLVED = "unresolved"


@dataclass(frozen=True, slots=True)
class FinalCall:
    """The reconciled outcome for one (firm, leaf) key across N run outputs."""

    firm: str
    firm_key: str
    leaf: str
    resolution: str
    unresolved: bool
    view: str                     # resolved view, or "" when unresolved
    views: tuple[str, ...]        # unresolved: the distinct differing views (both kept)
    commentary: str               # kept row's commentary (resolved); "" when unresolved
    note: str
    provenance: tuple[Row, ...]   # every contributing row


@dataclass(frozen=True, slots=True)
class _CCVerdict:
    bucket: str
    verdict: str
    note: str
    needs_human: bool


def _load_crosscheck(path: Path) -> dict[tuple[str, str], _CCVerdict]:
    """Load crosscheck.csv into {(firm_key, leaf): verdict}. Keyed identically to
    reconcile (imported `src.eval` normalization), so the join is exact."""
    if not path.is_file():
        raise SummarizeError(f"crosscheck file not found: {path}")
    verdicts: dict[tuple[str, str], _CCVerdict] = {}
    with path.open(newline="", encoding="utf-8-sig") as handle:
        for raw in csv.DictReader(handle):
            key = (normalize_firm(raw.get("Firm") or ""), _leaf_key(raw.get("Sub-Asset Class") or ""))
            verdicts[key] = _CCVerdict(
                bucket=(raw.get("bucket") or "").strip(),
                verdict=(raw.get("agent_verdict") or "").strip(),
                note=(raw.get("note") or "").strip(),
                needs_human=(raw.get("needs_human") or "").strip().lower() == "true",
            )
    return verdicts


def _confidence(row: Row) -> int:
    try:
        return int(row.confidence)
    except (TypeError, ValueError):
        return -1


def _date_key(row: Row) -> tuple[int, int, int]:
    """A sortable (Y, M, D) from an ``M/D/YYYY`` date; (0,0,0) when unparseable."""
    match = re.match(r"\s*(\d{1,2})/(\d{1,2})/(\d{4})", row.date or "")
    if not match:
        return (0, 0, 0)
    month, day, year = (int(match.group(i)) for i in (1, 2, 3))
    return (year, month, day)


def _keep_highest_confidence(rows: tuple[Row, ...]) -> Row:
    # Highest confidence wins; stable tie-break on (source_file, index) so the
    # kept row is deterministic across invocations.
    return min(rows, key=lambda r: (-_confidence(r), r.source_file, r.index))


def _keep_most_recent(rows: tuple[Row, ...]) -> Row:
    # "Superseded" == a more current reading wins: most recent date first, then
    # highest confidence, then the same stable tie-break.
    return min(
        rows,
        key=lambda r: (tuple(-v for v in _date_key(r)), -_confidence(r), r.source_file, r.index),
    )


def reconcile_firm_calls(
    output_paths: list[Path],
    crosscheck_path: Path | None = None,
) -> list[FinalCall]:
    """Reduce N run outputs (+ their cross-check) to one final call per key.

    Rules, per (firm, leaf):
      - one row              -> final as-is (`single`);
      - all-same-view rows   -> that view once, highest-confidence row kept
                                (`duplicate_same_view`);
      - conflicting + `superseded` -> the more-current row (most recent date,
                                then highest confidence) kept;
      - conflicting + `same_call`  -> the shared substance, highest-confidence
                                row kept once;
      - conflicting + `needs_human`/failed/absent verdict -> BOTH views kept,
                                flagged `unresolved` (never silently collapsed).
    A conflicting key with no crosscheck entry is `unresolved` by construction —
    we never pick a side without a verdict.

    The keep-rules are documented deterministic tie-breaks, NOT dual-confidence
    scoring; this is the v1 stopgap for the v1.2 firm-reconcile stage.
    """
    rows = load_rows(output_paths)
    crosscheck = _load_crosscheck(crosscheck_path) if crosscheck_path is not None else {}

    by_key: dict[tuple[str, str], list[Row]] = {}
    firm_display: dict[str, str] = {}
    for row in rows:
        by_key.setdefault((row.firm_key, row.leaf), []).append(row)
        firm_display.setdefault(row.firm_key, row.firm)

    finals: list[FinalCall] = []
    for (firm_key, leaf), members in by_key.items():
        ordered = tuple(sorted(members, key=lambda r: (r.source_file, r.index)))
        finals.append(_reconcile_key(firm_key, firm_display[firm_key], leaf, ordered, crosscheck))
    finals.sort(key=lambda f: (f.firm_key, f.leaf))
    return finals


def _reconcile_key(
    firm_key: str,
    firm: str,
    leaf: str,
    rows: tuple[Row, ...],
    crosscheck: dict[tuple[str, str], _CCVerdict],
) -> FinalCall:
    distinct_views = tuple(sorted({r.view for r in rows}))
    if len(rows) == 1:
        return _resolved(firm, firm_key, leaf, RES_SINGLE, rows[0], rows, "")
    if len(distinct_views) == 1:
        kept = _keep_highest_confidence(rows)
        note = f"{len(rows)} same-view rows across sources; kept highest-confidence provenance."
        return _resolved(firm, firm_key, leaf, RES_SAME_VIEW, kept, rows, note)

    cc = crosscheck.get((firm_key, leaf))
    if cc is not None and not cc.needs_human:
        if cc.verdict == "superseded":
            kept = _keep_most_recent(rows)
            return _resolved(firm, firm_key, leaf, RES_SUPERSEDED, kept, rows, cc.note)
        if cc.verdict == "same_call":
            kept = _keep_highest_confidence(rows)
            return _resolved(firm, firm_key, leaf, RES_SAME_CALL, kept, rows, cc.note)

    # needs_human, a failed/unknown verdict, or no crosscheck entry: keep both.
    note = cc.note if cc is not None and cc.note else "the firm's documents differ on this leaf; unresolved."
    return FinalCall(
        firm=firm,
        firm_key=firm_key,
        leaf=leaf,
        resolution=RES_UNRESOLVED,
        unresolved=True,
        view="",
        views=distinct_views,
        commentary="",
        note=note,
        provenance=rows,
    )


def _resolved(
    firm: str,
    firm_key: str,
    leaf: str,
    resolution: str,
    kept: Row,
    rows: tuple[Row, ...],
    note: str,
) -> FinalCall:
    return FinalCall(
        firm=firm,
        firm_key=firm_key,
        leaf=leaf,
        resolution=resolution,
        unresolved=False,
        view=kept.view,
        views=(kept.view,),
        commentary=kept.commentary,
        note=note,
        provenance=rows,
    )


# --------------------------------------------------------------------------- #
# Task 3 — Stage 2 per-firm page
# --------------------------------------------------------------------------- #


def parse_firm_page(raw_response: str) -> str:
    """Return the firm page as markdown; raise if it is missing its shape so the
    repair loop re-prompts. The page must open with an `# ` heading and carry a
    `## Sources` section (both required by the prompt)."""
    text = raw_response.strip()
    if text.startswith("```"):
        text = llm._extract_json(text)  # strips a stray code fence symmetrically
    if not text:
        raise ValueError("firm page is empty")
    if not any(line.startswith("# ") for line in text.splitlines()):
        raise ValueError("firm page must begin with an `# Firm` heading")
    if not re.search(r"(?im)^##\s+sources\b", text):
        raise ValueError("firm page must contain a `## Sources` section")
    return text


def _digest_sources(digests: list[dict]) -> list[dict[str, str]]:
    """Unique (title, url) source list from a firm's digests, title-sorted."""
    seen: dict[str, str] = {}
    for digest in digests:
        title = (digest.get("document_title") or "").strip()
        if title and title not in seen:
            seen[title] = (digest.get("url") or "").strip()
    return [{"title": title, "url": seen[title]} for title in sorted(seen)]


def _final_call_input(call: FinalCall) -> dict:
    payload: dict[str, object] = {
        "sub_asset_leaf": call.leaf,
        "resolution": call.resolution,
        "unresolved": call.unresolved,
    }
    if call.unresolved:
        payload["views"] = list(call.views)
        payload["divergence"] = [
            {
                "view": row.view,
                "source_title": row.source_title,
                "date": row.date,
                "commentary": row.commentary,
            }
            for row in call.provenance
        ]
    else:
        payload["view"] = call.view
        payload["commentary"] = call.commentary
    if call.note:
        payload["note"] = call.note
    return payload


def write_firm_page(
    firm: str,
    digests: list[dict],
    final_calls: list[FinalCall],
    *,
    engine: str,
    model: str | None,
    effort: str | None,
    runner: Runner | None = None,
) -> str:
    inputs = {
        "firm": firm,
        "digests": digests,
        "final_calls": [_final_call_input(call) for call in final_calls],
        "sources": _digest_sources(digests),
    }
    result = llm.call_parsed(
        FIRM_PAGE_PROMPT,
        inputs,
        engine=engine,
        model=model,
        effort=effort,
        runner=runner,
        parser=parse_firm_page,
    )
    return result.payload


def _load_digests(digests_dir: Path) -> dict[str, list[dict]]:
    """Load every digest JSON, grouped by normalized firm (title-sorted within)."""
    if not digests_dir.is_dir():
        raise SummarizeError(f"digests directory not found: {digests_dir}")
    by_firm: dict[str, list[dict]] = {}
    for path in sorted(digests_dir.glob("*.json")):
        digest = json.loads(path.read_text(encoding="utf-8"))
        firm_key = normalize_firm(digest.get("firm") or "")
        by_firm.setdefault(firm_key, []).append(digest)
    for digests in by_firm.values():
        digests.sort(key=lambda d: (d.get("document_title") or ""))
    return by_firm


def run_firmpages(
    digests_dir: Path,
    output_paths: list[Path],
    out_dir: Path,
    *,
    crosscheck_path: Path | None = None,
    firms: set[str] | None = None,
    engine: str = FIRMPAGE_ENGINE,
    model: str | None = FIRMPAGE_MODEL,
    effort: str | None = FIRMPAGE_EFFORT,
    runner: Runner | None = None,
) -> list[Path]:
    """Synthesize one markdown page per firm from its digests + reconciled calls.

    Fails loudly if a firm has multiple sources but no crosscheck file was given
    — conflicting rows must never be synthesized without verdicts. With a single
    source per selected firm and no crosscheck, reconciliation is pass-through.
    `firms` filters to a subset (matched on normalized firm name)."""
    by_firm = _load_digests(digests_dir)
    if firms is not None:
        wanted = {normalize_firm(f) for f in firms}
        by_firm = {k: v for k, v in by_firm.items() if k in wanted}
        missing = wanted - set(by_firm)
        if missing:
            raise SummarizeError(f"requested firms not found among digests: {sorted(missing)}")

    if crosscheck_path is None:
        multi = sorted(digest[0].get("firm", key) for key, digest in by_firm.items() if len(digest) > 1)
        if multi:
            raise SummarizeError(
                "these firms have multiple source digests but no --crosscheck was given; "
                f"refusing to synthesize conflicting rows without verdicts: {multi}"
            )

    finals = reconcile_firm_calls(output_paths, crosscheck_path)
    finals_by_firm: dict[str, list[FinalCall]] = {}
    for call in finals:
        finals_by_firm.setdefault(call.firm_key, []).append(call)

    out_dir.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []
    for firm_key in sorted(by_firm):
        digests = by_firm[firm_key]
        firm = digests[0].get("firm") or firm_key
        page = write_firm_page(
            firm,
            digests,
            finals_by_firm.get(firm_key, []),
            engine=engine,
            model=model,
            effort=effort,
            runner=runner,
        )
        path = out_dir / f"{slugify(firm)}.md"
        path.write_text(page.rstrip() + "\n", encoding="utf-8")
        written.append(path)
    return written


# --------------------------------------------------------------------------- #
# Task 4 — deterministic python-docx binder (no LLM)
# --------------------------------------------------------------------------- #

_MD_LINK = re.compile(r"\[([^\]]+)\]\(([^)]*)\)")
_BOLD = re.compile(r"\*\*([^*]+)\*\*")


def _strip_inline(text: str) -> str:
    return _BOLD.sub(r"\1", text).strip()


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Add a hyperlink run to a python-docx paragraph. Falls back to visible
    ``text (url)`` plain text when no URL is present."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not url:
        paragraph.add_run(text)
        return
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    style = OxmlElement("w:rStyle")
    style.set(qn("w:val"), "Hyperlink")
    run_props.append(style)
    run.append(run_props)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    link.append(run)
    paragraph._p.append(link)


def _render_page(document, markdown: str) -> None:
    """Render one firm markdown page into the docx (headings, bullets, sources)."""
    in_sources = False
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            document.add_heading(_strip_inline(line[2:]), level=1)
            in_sources = False
        elif line.startswith("## "):
            heading = _strip_inline(line[3:])
            document.add_heading(heading, level=2)
            in_sources = heading.lower().startswith("sources")
        elif line.startswith(("- ", "* ")):
            _render_bullet(document, line[2:].strip(), in_sources)
        elif re.match(r"^\d+\.\s+", line):
            document.add_paragraph(_strip_inline(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
        else:
            document.add_paragraph(_strip_inline(line))


def _render_bullet(document, body: str, in_sources: bool) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    link = _MD_LINK.search(body)
    if in_sources and link:
        _add_hyperlink(paragraph, link.group(2).strip(), link.group(1).strip())
    else:
        paragraph.add_run(_strip_inline(body))


def bind_pages(pages_dir: Path, out_path: Path, *, title: str | None = None) -> Path:
    """Merge every ``*.md`` firm page (filename-sorted) into one Word document,
    a page break between firms. Content is deterministic given the same pages."""
    from docx import Document

    pages = sorted(pages_dir.glob("*.md"))
    if not pages:
        raise SummarizeError(f"no firm pages (*.md) found in {pages_dir}")

    document = Document()
    if title:
        document.add_heading(title, level=0)
        document.add_page_break()
    for index, page in enumerate(pages):
        if index or title:
            document.add_page_break()
        _render_page(document, page.read_text(encoding="utf-8"))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    return out_path


# --------------------------------------------------------------------------- #
# CLI
# --------------------------------------------------------------------------- #


def _resolve_model(engine: str, model: str | None, default_model: str | None) -> str | None:
    """codex is pinned to its own model; claude needs an explicit model, so fill
    the stage default when the flag is omitted."""
    if engine == "codex":
        return None
    return model if model is not None else default_model


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        prog="python -m src.summarize",
        description="Reader summaries: per-source digests -> per-firm pages -> Word binder.",
    )
    sub = parser.add_subparsers(dest="command", required=True)

    p_digest = sub.add_parser("digest", help="Stage 1: per-source digests from a frozen run")
    p_digest.add_argument("--run", required=True, type=Path, help="frozen run dir (runs/<name>; read-only)")
    p_digest.add_argument("--out-dir", required=True, type=Path)
    p_digest.add_argument("--sources", nargs="+", default=None, help="subset of source ids (default: all)")
    p_digest.add_argument("--engine", default=DIGEST_ENGINE)
    p_digest.add_argument("--model", default=None)
    p_digest.add_argument("--effort", default=DIGEST_EFFORT)

    p_firm = sub.add_parser("firmpages", help="Stage 2: per-firm markdown pages")
    p_firm.add_argument("--digests", required=True, type=Path)
    p_firm.add_argument("--outputs", required=True, nargs="+", type=Path)
    p_firm.add_argument("--crosscheck", default=None, type=Path)
    p_firm.add_argument("--out-dir", required=True, type=Path)
    p_firm.add_argument("--firms", nargs="+", default=None)
    p_firm.add_argument("--engine", default=FIRMPAGE_ENGINE)
    p_firm.add_argument("--model", default=None)
    p_firm.add_argument("--effort", default=FIRMPAGE_EFFORT)

    p_bind = sub.add_parser("bind", help="deterministic python-docx binder")
    p_bind.add_argument("--pages", required=True, type=Path)
    p_bind.add_argument("--out", required=True, type=Path)
    p_bind.add_argument("--title", default=None)

    args = parser.parse_args(argv)

    if args.command == "digest":
        model = _resolve_model(args.engine, args.model, DIGEST_MODEL)
        written = run_digests(
            args.run,
            args.out_dir,
            only=set(args.sources) if args.sources else None,
            engine=args.engine,
            model=model,
            effort=args.effort,
        )
        print(f"digest: wrote {len(written)} source digests -> {args.out_dir}")
    elif args.command == "firmpages":
        model = _resolve_model(args.engine, args.model, FIRMPAGE_MODEL)
        written = run_firmpages(
            args.digests,
            args.outputs,
            args.out_dir,
            crosscheck_path=args.crosscheck,
            firms=set(args.firms) if args.firms else None,
            engine=args.engine,
            model=model,
            effort=args.effort,
        )
        print(f"firmpages: wrote {len(written)} firm pages -> {args.out_dir}")
    elif args.command == "bind":
        out = bind_pages(args.pages, args.out, title=args.title)
        print(f"bind: wrote {out}")
    return 0


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())
